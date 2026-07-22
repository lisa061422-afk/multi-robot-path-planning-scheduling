from __future__ import annotations

import os
import sys
from datetime import date
from pathlib import Path

DEPS = Path(__file__).resolve().parent / ".docx_deps"
sys.path.insert(0, str(DEPS))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("机器人路径规划与Sim-to-Real实验路线参考.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "1F2933"
CAUTION = "7A5A00"

LATIN_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    assert sum(widths_dxa) == 9360, widths_dxa
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[i] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, *, size=None, bold=None, italic=None, color=BLACK, font=LATIN_FONT) -> None:
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_fonts.set(qn("w:cs"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph_runs(paragraph, size=10.5, color=BLACK, bold=False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def set_cell_text(cell, text: str, *, bold=False, color=BLACK, size=9.3,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = LATIN_FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Lead" not in styles:
        st = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles["Lead"]
    st.font.name = LATIN_FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    st.font.size = Pt(12)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(INK)
    st.paragraph_format.space_before = Pt(2)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        st = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles["Code Block"]
    st.font.name = "Consolas"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    st._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    st.font.size = Pt(9.2)
    st.font.color.rgb = RGBColor.from_string(INK)
    st.paragraph_format.left_indent = Inches(0.22)
    st.paragraph_format.right_indent = Inches(0.12)
    st.paragraph_format.space_before = Pt(4)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.12
    p_pr = st.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p_pr.append(shd)


def create_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element

    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract, default=-1) + 1

    def add_abstract(abs_id: int, num_fmt: str, text: str, font: str | None = None) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, lvl_text, jc, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    add_abstract(abstract_id, "bullet", "•", "Symbol")
    add_abstract(abstract_id + 1, "decimal", "%1.")

    existing_nums = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_num = max(existing_nums, default=0) + 1

    def add_num(num_id: int, abs_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)

    add_num(next_num, abstract_id)
    add_num(next_num + 1, abstract_id + 1)
    return next_num, next_num + 1


def add_list_item(doc: Document, text: str, num_id: int, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=BLACK)


def add_callout(doc: Document, label: str, text: str, fill=CALLOUT, accent=BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              *, font_size=9.1) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    repeat_header(hdr)
    prevent_row_split(hdr)
    for i, h in enumerate(headers):
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        set_cell_text(hdr.cells[i], h, bold=True, color=INK, size=9.2,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=font_size,
                          align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, bold=True, color=INK)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=BLACK)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size=9.2, color=INK, font="Consolas")


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def configure_header_footer(section) -> None:
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("机器人路径规划与 Sim-to-Real 实验路线参考")
    set_run_font(r, size=8.7, color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("参考指南  |  ")
    set_run_font(r, size=8.7, color=MUTED)
    add_page_field(fp)


def build_document() -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    bullet_id, decimal_id = create_numbering(doc)
    section = doc.sections[0]
    configure_header_footer(section)
    section.different_first_page_header_footer = True

    # Cover: editorial_cover pattern, compact-reference palette.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ROBOTICS PROJECT GUIDE")
    set_run_font(r, size=10.5, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("机器人路径规划与\nSim-to-Real 实验路线参考")
    set_run_font(r, size=27, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于多车路网路径选择、交叉口调度与毕业前真机验证的实施建议")
    set_run_font(r, size=13.5, color=DARK_BLUE)

    add_callout(
        doc,
        "核心判断",
        "当前项目属于离散图上的多车路径选择与资源调度联合优化。毕业前最稳妥的真机路线，是三辆缩比 Ackermann 小车、顶视相机闭环定位和中央调度，而不是直接扩展到三台机器狗。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("个人项目规划参考")
    set_run_font(r, size=11, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年7月")
    set_run_font(r, size=10, color=MUTED)

    doc.add_page_break()

    add_heading(doc, "阅读导航", 1)
    for t in [
        "一页结论：项目如何命名、毕业前应做什么",
        "当前算法的准确定位与学术边界",
        "Route、Path、Motion 与 Trajectory Planning 的区别",
        "Motion Planning 岗位需要的技能",
        "适合当前项目的 Sim-to-Real 硬件架构",
        "毕业前 12 周真机实施计划与实验指标",
        "机器狗方向：学习内容、技术栈与取舍",
        "最终建议、简历表述与下一步清单",
    ]:
        add_list_item(doc, t, bullet_id)

    add_heading(doc, "1. 一页结论", 1)
    p = doc.add_paragraph(style="Lead")
    p.add_run("项目可以作为 path planning 相关项目展示，但最准确的技术标签是 graph-based route selection and intersection scheduling co-design。")
    style_paragraph_runs(p, size=12, color=INK, bold=True)

    add_table(doc,
              ["问题", "建议"],
              [
                  ["求职标题", "Graph-Based Multi-Vehicle Path Planning and Intersection Scheduling"],
                  ["学术定位", "离散图上的候选路径选择、资源冲突约束与精确组合搜索"],
                  ["避免过度包装", "暂不单独称为 Real-Time Dynamic Motion Planner 或 Autonomous Driving Motion Planner"],
                  ["毕业前真机", "3辆缩比 Ackermann 小车 + 顶视相机 + 中央计算机 + 每车低层控制"],
                  ["机器狗", "若实验室已有设备，可作为异构机器人扩展；不建议替代主线真机平台"],
              ],
              [2200, 7160], font_size=9.4)

    add_callout(doc, "优先级", "可靠完成三车闭环实验 > 增加昂贵传感器 > 追求复杂车辆动力学或足式控制。")

    add_heading(doc, "推荐项目名称", 2)
    for item in [
        "求职通用：Graph-Based Multi-Vehicle Path Planning and Intersection Scheduling",
        "算法/运筹方向：Joint Route Selection and Conflict-Free Intersection Scheduling via Branch-and-Bound",
        "学术论文风格：An Exact Graph-Based Co-Design Method for Multi-Vehicle Routing and Intersection Scheduling",
        "完成真机后：Sim-to-Real Multi-Vehicle Route and Motion Coordination with Closed-Loop Intersection Scheduling",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "2. 当前算法的准确定位", 1)
    add_body(doc, "当前程序先为每个起终点对枚举所有 simple intersection paths，再在联合搜索过程中依据车辆的任务生成时刻和路径前缀逐步缩小候选路线，同时搜索交叉口资源占用顺序。目标函数为：")
    add_code_block(doc, "J = total intersection delay + path_extra")

    add_table(doc,
              ["维度", "当前实现", "准确解读"],
              [
                  ["空间表示", "交叉口与道路组成的离散图", "属于 graph-based planning，而非连续几何空间规划"],
                  ["路径候选", "默认枚举 OD 间全部 simple paths", "小规模网络可作为 exact method；规模扩大时可能指数增长"],
                  ["动态决策", "到下一任务/分叉点时按路线前缀分支", "属于 deferred route commitment"],
                  ["调度", "每个交叉口作为互斥资源", "与 job-shop/resource-constrained scheduling 接近"],
                  ["求解", "DFS + branch-and-bound，可并行", "属于精确组合优化与最优解搜索"],
                  ["输出", "最佳路径组合与交叉口时间安排", "高层 route/schedule，不是可直接执行的连续轨迹"],
              ],
              [1700, 3430, 4230], font_size=8.8)

    add_heading(doc, "学术上是否属于标准 Path Planning", 2)
    add_body(doc, "它属于 path planning/routing 的正规子问题，尤其适合描述为 path-based formulation、candidate-route enumeration 或 routing-and-scheduling co-design。对于小规模网络，穷举候选路线配合 branch-and-bound 可以提供最优性基准。")
    add_body(doc, "但它不是机器人领域通用的大规模 motion-planning 标准算法。所有 simple paths 的数量可能随网络规模指数增长；而且当前程序在已知全部车辆请求后先算出完整最优解，再执行最佳方案，因此严格意义上仍是 offline optimization with deferred route commitment。")

    add_heading(doc, "何时可以称为 Online/Dynamic Planner", 2)
    for item in [
        "车辆执行过程中持续接收实际位置、速度和路口占用状态；",
        "新车辆、道路阻塞或时间偏差出现后可以更新问题；",
        "以滚动时域或事件触发方式重新规划；",
        "在明确的实时预算内返回可执行且安全的结果。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "3. Route、Path、Motion 与 Trajectory Planning", 1)
    add_table(doc,
              ["层次", "核心问题", "典型输出", "当前项目覆盖"],
              [
                  ["Route planning", "经过哪些道路/路口", "I1 → I4 → I3", "强"],
                  ["Path planning", "在空间中沿哪条几何曲线", "x(s), y(s), theta(s)", "部分"],
                  ["Trajectory planning", "每个时刻的位置与速度", "x(t), y(t), theta(t), v(t)", "尚未"],
                  ["Motion planning", "满足碰撞、运动学、动力学与时间约束的运动", "状态和控制随时间变化", "尚未完整覆盖"],
                  ["Control", "如何让执行器跟踪轨迹", "转向角、加速度、关节力矩", "尚未"],
              ],
              [1700, 3000, 2800, 1860], font_size=8.7)

    add_callout(doc, "关键区别", "Path 通常描述几何曲线；trajectory 显式包含时间。Motion planning 是更宽的概念，在招聘语境中经常与 trajectory planning 混用。")

    add_heading(doc, "4. Motion Planning 岗位需要的技能", 1)
    add_heading(doc, "4.1 建模与基础理论", 2)
    for item in [
        "状态与控制建模：位置、姿态、速度、加速度、转向角；",
        "Kinematic bicycle model、dynamic bicycle model 或 differential-drive model；",
        "坐标系、刚体运动、曲率与非完整约束；",
        "离散时间系统、数值积分与约束可达性。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "4.2 搜索、优化与碰撞", 2)
    add_table(doc,
              ["技能组", "典型内容"],
              [
                  ["几何搜索", "A*、Dijkstra、Hybrid A*、state lattice、RRT/RRT*、PRM"],
                  ["轨迹优化", "QP、NLP、SQP、direct collocation、spline optimization、MPC"],
                  ["碰撞检测", "Occupancy grid、polygon/OBB、configuration space、space-time collision"],
                  ["动态可行性", "速度、加速度、转向角、曲率、横向加速度和 jerk 约束"],
                  ["实时系统", "Receding horizon、deadline、fallback trajectory、emergency stop"],
              ],
              [2200, 7160], font_size=9.1)

    add_heading(doc, "4.3 与当前项目的最佳集成方式", 2)
    add_code_block(doc,
                   "全局路线与交叉口调度器\n"
                   "        ↓ route + intersection time windows\n"
                   "道路中心线 / Hybrid A* / state lattice\n"
                   "        ↓ geometric path\n"
                   "轨迹优化或速度规划\n"
                   "        ↓ x(t), y(t), theta(t), v(t)\n"
                   "Pure Pursuit / Stanley / MPC tracking controller")
    add_body(doc, "这种分层设计保留现有调度研究的独特性，同时补上 motion planning 求职中最重要的连续轨迹、车辆约束和闭环执行能力。")

    add_heading(doc, "5. Sim-to-Real 真机平台建议", 1)
    p = doc.add_paragraph(style="Lead")
    p.add_run("推荐平台：3辆相同的 1/16 至 1/10 Ackermann 小车，顶视相机集中定位，中央电脑运行 planner，每辆车只负责轨迹跟踪与底层控制。")
    style_paragraph_runs(p, size=12, color=INK, bold=True)

    add_heading(doc, "5.1 系统链路", 2)
    add_code_block(doc,
                   "顶视相机 + AprilTag\n"
                   "        ↓ 全局 x, y, theta\n"
                   "中央电脑：route selection + scheduling\n"
                   "        ↓ Wi-Fi / ROS 2：route、time window、velocity command\n"
                   "每车 Raspberry Pi / SBC\n"
                   "        ↓ steering、speed reference\n"
                   "MCU → ESC/电机 + 转向舵机\n"
                   "        ↑ encoder + IMU")

    add_heading(doc, "5.2 每辆车的设备", 2)
    for item in [
        "Ackermann 转向 RC chassis，优先选择可替换标准舵机和 ESC 的型号；",
        "转向舵机、电机与 ESC；",
        "Raspberry Pi 5 或类似 SBC，负责 ROS 2、通信、状态融合与轨迹跟踪；",
        "ESP32、STM32 或 Teensy，负责 PWM、电机和实时安全逻辑；",
        "轮速编码器、IMU；",
        "电池、保险丝、合适的 DC-DC 电源模块；",
        "物理急停或遥控 kill switch；",
        "车顶唯一 AprilTag，前后方向必须清晰可辨。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "5.3 共享设备", 2)
    for item in [
        "现有电脑：运行中央调度器、可视化、日志和实验统计；",
        "独立 Wi-Fi 路由器，避免使用校园或家庭拥堵网络；",
        "顶视 global-shutter USB 相机和稳定支架；",
        "打印或胶带铺设的 2×2/3×3 路网、泡沫护栏和停止线；",
        "LiPo 充电器、防火袋、备用电池、备用舵机和轮胎。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "5.4 计算平台怎么选", 2)
    add_table(doc,
              ["方案", "适合场景", "建议"],
              [
                  ["Raspberry Pi 5", "通信、状态估计、轨迹跟踪、轻量规划", "当前项目首选；中央电脑承担组合搜索"],
                  ["Jetson Orin Nano", "车载神经网络、深度视觉、端到端感知", "只有加入 onboard AI perception 时再购买"],
                  ["完整 F1TENTH", "自动驾驶研究、LiDAR、标准化 1/10 平台", "求职展示强，但三车成本和维护负担较高"],
                  ["顶视相机集中定位", "固定实验场地、多车协调验证", "最适合毕业前快速完成可靠实验"],
              ],
              [1900, 3500, 3960], font_size=8.9)

    add_callout(doc, "采购策略", "先买一辆车打通全链路，再购买另外两辆同型号车。同步准备备用电池、舵机和线材，避免实验阶段因小部件停工。")

    add_heading(doc, "5.5 安全与实验边界", 2)
    for item in [
        "第一阶段限制最高速度，并使用泡沫护栏和空旷场地；",
        "定位超时、通信中断、车辆超出地图或错过时间窗口时默认停车；",
        "急停链路必须独立于上层 planner；",
        "电池和充电遵循厂家要求，不在无人看管状态下充电；",
        "任何动态重规划失败都应回退到停止，而不是继续执行过期轨迹。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "6. 毕业前 12 周实施计划", 1)
    add_table(doc,
              ["时间", "里程碑", "主要工作", "验收标准"],
              [
                  ["第1–2周", "单车基础", "底盘、遥控、ROS通信、转向和速度控制", "可重复发送速度/转向命令；急停有效"],
                  ["第3–4周", "全局定位", "相机标定、AprilTag、坐标变换和时间戳", "位置姿态连续稳定；短时遮挡可安全停车"],
                  ["第5–6周", "单车闭环", "道路中心线、Pure Pursuit/Stanley、速度 PID", "单车完成指定 route；误差可量化"],
                  ["第7–8周", "两车协调", "停止线、时间窗口、冲突区放行逻辑", "多次通过同一路口且无碰撞"],
                  ["第9–10周", "三车联合优化", "接入 route selection 和 schedule；处理早到/晚到", "展示非最短路线降低系统总延误"],
                  ["第11周", "批量实验", "baseline、重复试验、日志与统计", "每种方法至少10–20次有效实验"],
                  ["第12周", "交付", "稳定性修复、图表、视频、README和简历素材", "可一键复现实验并解释全部指标"],
              ],
              [1250, 1850, 3380, 2880], font_size=8.25)

    add_heading(doc, "7. 真机实验设计", 1)
    add_heading(doc, "7.1 必须对比的 Baselines", 2)
    for item in [
        "Fixed shortest paths + intersection scheduling；",
        "Joint route selection + intersection scheduling；",
        "理想仿真执行；",
        "实际车辆闭环执行。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "7.2 建议记录的指标", 2)
    add_table(doc,
              ["指标组", "具体指标", "用于回答的问题"],
              [
                  ["效率", "total delay、makespan、travel distance", "绕路是否换来了系统级效率提升？"],
                  ["求解", "planning time、expanded/pruned nodes", "算法是否能在目标规模内及时返回？"],
                  ["执行", "arrival-time error、lateral tracking error", "仿真 schedule 能否被真车准确执行？"],
                  ["安全", "minimum distance、emergency stops、collisions", "误差和延迟下是否仍保持安全？"],
                  ["稳健性", "成功率、重复试验方差、通信/定位丢失次数", "结果是否可重复，而不是一次性演示？"],
              ],
              [1700, 3300, 4360], font_size=8.8)

    add_heading(doc, "7.3 执行监控器", 2)
    add_code_block(doc,
                   "当前位置 + 当前时间 + 下一路口 time window\n"
                   "        ↓\n"
                   "能按时到达？ ── 是 → 正常轨迹跟踪\n"
                   "        ├─ 太早 → 减速或入口等待\n"
                   "        └─ 太晚/偏差过大 → 安全停车或重新调度")
    add_body(doc, "第一版可以采用安全停车作为 fallback；当单车和三车执行稳定后，再加入事件触发或 receding-horizon rescheduling。")

    add_heading(doc, "8. 机器狗方向：它学习的是什么", 1)
    add_body(doc, "机器狗属于 legged robotics。核心问题是在足端不断接触和离地的混合动力学过程中，协调多个关节、维持身体稳定并跟踪期望速度。它比轮式车多了 gait、footstep、contact dynamics 和 whole-body control。")

    add_table(doc,
              ["层次", "解决的问题", "典型方法"],
              [
                  ["Global navigation", "整体往哪里走", "A*、Nav2、地图规划"],
                  ["Footstep planning", "每只脚下一步踩在哪里", "搜索、优化、terrain-aware planning"],
                  ["Locomotion", "采用什么步态、身体如何推进", "gait generator、MPC、RL policy"],
                  ["Whole-body control", "身体和四条腿如何协调", "QP/WBC、contact constraints"],
                  ["Joint control", "关节输出位置、速度或力矩", "PD、torque control"],
                  ["State estimation", "估计机身、速度与接触状态", "IMU、encoder、contact、Kalman filter"],
              ],
              [1900, 3500, 3960], font_size=8.8)

    add_heading(doc, "8.1 强化学习 Policy 的输入与输出", 2)
    add_table(doc,
              ["组成", "常见内容"],
              [
                  ["Observation", "IMU、身体角速度、关节角/速度、足端接触、上一动作、期望线速度/角速度、可选地形高度"],
                  ["Action", "关节目标角度、位置增量或直接关节力矩"],
                  ["Reward", "速度跟踪、身体直立、足端抬高；惩罚滑动、能耗、jerk和跌倒"],
                  ["Sim-to-real", "随机化质量、摩擦、电机强度、传感器噪声、时延和地形"],
              ],
              [2200, 7160], font_size=9.0)

    add_heading(doc, "8.2 商业机器狗 SDK 与真正的 Legged Control", 2)
    add_body(doc, "如果只通过商业机器狗 SDK 发送 vx、vy 和 yaw_rate，厂家已经完成步态和关节控制。此时你的贡献仍然是高层 multi-robot routing and scheduling，不能把项目表述为自研 legged locomotion。")
    add_body(doc, "只有自己实现或显著改进 RL locomotion、footstep planner、MPC/WBC、joint-level control、domain randomization 和真机部署，才属于完整的足式机器人运动控制工作。")

    add_heading(doc, "8.3 是否应把毕业真机改成机器狗", 2)
    add_table(doc,
              ["条件", "建议"],
              [
                  ["目标是按时验证当前算法", "继续使用轮式 Ackermann 小车；变量少、成本低、三车更容易稳定"],
                  ["实验室已有机器狗且可使用 SDK", "可将一台作为异构机器人扩展，不替代三车主实验"],
                  ["目标转向 legged-robotics 求职", "另开 locomotion/MPC/RL 项目，并确认设备开放低层关节或力矩接口"],
                  ["需要视觉展示", "机器狗很吸引眼球，但应清楚区分平台能力与自己的算法贡献"],
              ],
              [3000, 6360], font_size=9.0)

    add_callout(doc, "推荐组合", "先用三辆轮式车完成严谨、可重复的主实验；若实验室已有机器狗，再把其中一台接入同一调度系统，扩展为 heterogeneous multi-robot coordination。")

    add_heading(doc, "9. 最终实施建议", 1)
    add_heading(doc, "阶段 A：毕业前必须完成", 2)
    for item in [
        "把 scheduler 输出转换为 route、waypoints 和 intersection time windows；",
        "实现顶视定位、单车闭环跟踪和独立安全停车；",
        "完成两车冲突验证和三车路径—调度联合实验；",
        "建立 baseline、重复实验、统计结果和完整视频；",
        "将实验配置、标定步骤和复现命令写入 README。",
    ]:
        add_list_item(doc, item, decimal_id)

    add_heading(doc, "阶段 B：时间允许再增加", 2)
    for item in [
        "事件触发或滚动时域重新调度；",
        "Hybrid A* 或 state lattice 几何规划；",
        "带速度、加速度和 jerk 限制的轨迹优化；",
        "LiDAR/深度相机和动态障碍物；",
        "异构机器人或机器狗接入。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "10. 求职表述模板", 1)
    add_heading(doc, "当前阶段可以使用", 2)
    add_callout(doc, "Project title", "Graph-Based Multi-Vehicle Path Planning and Intersection Scheduling")
    for item in [
        "Developed an exact graph-based co-design solver that jointly optimizes discrete route choices and conflict-free intersection schedules.",
        "Implemented prefix-level dynamic route branching with DFS branch-and-bound and optional parallel frontier search.",
        "Evaluated shortest-route baselines against delay-plus-path-cost optimization on multi-intersection traffic networks.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "完成真机后再使用", 2)
    for item in [
        "Deployed the planner on three scaled Ackermann vehicles using overhead visual localization and ROS 2 closed-loop trajectory tracking.",
        "Designed time-window execution and fail-safe supervision to handle localization error, communication delay, and schedule deviation.",
        "Validated sim-to-real performance through repeated baseline experiments measuring delay, makespan, tracking error, and minimum separation.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_callout(doc, "表述原则", "没有完成真机前，不把 planned work 写成 deployed/validated；没有在线更新外部状态前，不单独声称 real-time online replanning。")

    add_heading(doc, "11. 下一步清单", 1)
    for item in [
        "确定毕业前剩余时间与总预算；",
        "确定实验场地尺寸、相机安装高度和可用计算机；",
        "选择一辆可编程 Ackermann 小车并先完成单车原型；",
        "定义 scheduler-to-controller 接口消息；",
        "实现路线中心线、时间窗口与安全状态机；",
        "建立仿真/真机统一日志格式；",
        "在单车验收后复制到三辆同型号车辆；",
        "提前定义实验表格、指标和视频脚本。",
    ]:
        add_list_item(doc, item, decimal_id)

    add_heading(doc, "参考资料", 1)
    sources = [
        ("当前项目源码：main.py、traffic_map.py、coarse_scheduler.py", "本地工作区"),
        ("F1TENTH 官方构建与平台资料", "https://f1tenth.org/build.html"),
        ("F1TENTH 1/10 Reference Manual", "https://f1tenth.org/build/BuildV2.pdf"),
        ("Raspberry Pi 计算机硬件文档", "https://www.raspberrypi.com/documentation/computers/raspberry-pi.html"),
        ("NVIDIA Jetson FAQ", "https://developer.nvidia.com/embedded/faq"),
        ("ROS 2 Navigation / Nav2 文档", "https://docs.nav2.org/"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.space_after = Pt(4)
        if url.startswith("http"):
            add_hyperlink(p, label, url)
        else:
            r = p.add_run(label)
            set_run_font(r, size=10.3, color=BLACK)
        r = p.add_run(f" — {url}")
        set_run_font(r, size=9.5, color=MUTED)

    add_body(doc, "说明：设备型号、价格和供货状态会变化。采购前应再次核对厂家接口、电源要求、ROS 2 支持和低层控制开放程度。")

    # Set metadata and ensure all sections use the preset page geometry.
    doc.core_properties.title = "机器人路径规划与 Sim-to-Real 实验路线参考"
    doc.core_properties.subject = "多车路径规划、运动规划、真机实验与机器狗方向建议"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "path planning, motion planning, sim-to-real, multi-vehicle, legged robotics"
    configure_page(doc)
    for sec in doc.sections:
        if sec is not doc.sections[0]:
            configure_header_footer(sec)

    doc.save(OUT)
    print("DOCX created successfully")


if __name__ == "__main__":
    build_document()
