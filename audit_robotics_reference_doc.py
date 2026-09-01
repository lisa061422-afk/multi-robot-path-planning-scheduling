from pathlib import Path
import sys

DEPS = Path(__file__).resolve().parent / ".docx_deps"
sys.path.insert(0, str(DEPS))

from docx import Document
from docx.oxml.ns import qn


path = Path("机器人路径规划与Sim-to-Real实验路线参考.docx")
doc = Document(path)
errors = []

if len(doc.sections) != 1:
    errors.append(f"expected 1 section, got {len(doc.sections)}")

section = doc.sections[0]
expected = {
    "page_width": 8.5,
    "page_height": 11.0,
    "top_margin": 1.0,
    "bottom_margin": 1.0,
    "left_margin": 1.0,
    "right_margin": 1.0,
}
for attr, target in expected.items():
    actual = getattr(section, attr).inches
    if abs(actual - target) > 0.01:
        errors.append(f"{attr}: expected {target}, got {actual}")

heading_count = 0
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading"):
        heading_count += 1
        if not p.text.strip():
            errors.append(f"empty heading at paragraph {i}")
    if p.text.startswith(("•", "- ")):
        errors.append(f"fake list marker at paragraph {i}: {p.text[:50]}")

for ti, table in enumerate(doc.tables):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_w is None or tbl_w.get(qn("w:w")) != "9360" or tbl_w.get(qn("w:type")) != "dxa":
        errors.append(f"table {ti}: tblW mismatch")
    if tbl_ind is None or tbl_ind.get(qn("w:w")) != "120":
        errors.append(f"table {ti}: tblInd mismatch")
    widths = [int(c.get(qn("w:w"))) for c in table._tbl.tblGrid.findall(qn("w:gridCol"))]
    if sum(widths) != 9360:
        errors.append(f"table {ti}: grid width sums to {sum(widths)}")
    first_tr_pr = table.rows[0]._tr.trPr
    if first_tr_pr is None or first_tr_pr.find(qn("w:tblHeader")) is None:
        errors.append(f"table {ti}: first row not marked header")
    for ri, row in enumerate(table.rows):
        if len(row.cells) != len(widths):
            errors.append(f"table {ti} row {ri}: cell count mismatch")
            continue
        for ci, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None or int(tc_w.get(qn("w:w"))) != widths[ci]:
                errors.append(f"table {ti} row {ri} col {ci}: tcW mismatch")

print(f"paragraphs={len(doc.paragraphs)}")
print(f"headings={heading_count}")
print(f"tables={len(doc.tables)}")
print(f"file_bytes={path.stat().st_size}")
if errors:
    print("AUDIT FAILED")
    for error in errors:
        print(error)
    raise SystemExit(1)
print("AUDIT PASSED")
